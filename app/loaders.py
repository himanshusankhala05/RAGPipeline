import csv
import hashlib
import json
from pathlib import Path
from typing import Any

from docx import Document
from langchain_text_splitters import (
	CharacterTextSplitter,
	MarkdownHeaderTextSplitter,
	PythonCodeTextSplitter,
	RecursiveCharacterTextSplitter,
	TextSplitter,
	TokenTextSplitter,
)
from openpyxl import load_workbook
from pypdf import PdfReader


SUPPORTED_TEXT_EXTENSIONS = {
	".csv",
	".html",
	".json",
	".md",
	".py",
	".txt",
	".xml",
}

CHUNKING_STRATEGIES = {
	"Recursive character": "recursive_character",
	"Character": "character",
	"Token": "token",
	"Markdown header": "markdown_header",
	"Python code": "python_code",
	"Semantic": "semantic",
}


def _read_text_file(file_path: Path) -> str:
	if file_path.suffix.lower() == ".csv":
		with file_path.open(newline="", encoding="utf-8-sig") as file:
			rows = csv.reader(file)
			return "\n".join(" | ".join(row) for row in rows)

	if file_path.suffix.lower() == ".json":
		data = json.loads(file_path.read_text(encoding="utf-8"))
		return json.dumps(data, indent=2, ensure_ascii=False)

	return file_path.read_text(encoding="utf-8-sig")


def _read_pdf_file(file_path: Path) -> str:
	reader = PdfReader(str(file_path))
	pages = [page.extract_text() or "" for page in reader.pages]
	return "\n\n".join(pages)


def _read_docx_file(file_path: Path) -> str:
	document = Document(str(file_path))
	paragraphs = [paragraph.text for paragraph in document.paragraphs]
	table_rows = [
		" | ".join(cell.text for cell in row.cells)
		for table in document.tables
		for row in table.rows
	]
	return "\n".join(paragraphs + table_rows)


def _read_xlsx_file(file_path: Path) -> str:
	workbook = load_workbook(filename=file_path, read_only=True, data_only=True)
	rows: list[str] = []

	for worksheet in workbook.worksheets:
		rows.append(f"Sheet: {worksheet.title}")
		for row in worksheet.iter_rows(values_only=True):
			values = [str(value) for value in row if value is not None]
			if values:
				rows.append(" | ".join(values))

	workbook.close()
	return "\n".join(rows)


def extract_text(file_path: str | Path) -> str:
	"""Extract searchable text from a supported file."""
	path = Path(file_path)
	extension = path.suffix.lower()

	if extension in SUPPORTED_TEXT_EXTENSIONS:
		text = _read_text_file(path)
	elif extension == ".pdf":
		text = _read_pdf_file(path)
	elif extension == ".docx":
		text = _read_docx_file(path)
	elif extension == ".xlsx":
		text = _read_xlsx_file(path)
	else:
		supported = ", ".join(sorted(SUPPORTED_TEXT_EXTENSIONS | {".docx", ".pdf", ".xlsx"}))
		raise ValueError(
			f"Unsupported file type '{extension}'. Supported types: {supported}"
		)

	text = text.strip()
	if not text:
		raise ValueError(f"No readable text found in '{path.name}'")
	return text


def _create_splitter(
	strategy: str,
	chunk_size: int,
	chunk_overlap: int,
) -> TextSplitter:
	if strategy == "recursive_character":
		return RecursiveCharacterTextSplitter(
			chunk_size=chunk_size,
			chunk_overlap=chunk_overlap,
		)

	if strategy == "character":
		return CharacterTextSplitter(
			separator="\n\n",
			chunk_size=chunk_size,
			chunk_overlap=chunk_overlap,
		)

	if strategy == "token":
		return TokenTextSplitter(
			chunk_size=chunk_size,
			chunk_overlap=chunk_overlap,
		)

	if strategy == "python_code":
		return PythonCodeTextSplitter(
			chunk_size=chunk_size,
			chunk_overlap=chunk_overlap,
		)

	if strategy == "markdown_header":
		return MarkdownHeaderTextSplitter(
			headers_to_split_on=[
				("#", "Header 1"),
				("##", "Header 2"),
				("###", "Header 3"),
			],
			strip_headers=False,
		)

	if strategy == "semantic":
		from langchain_experimental.text_splitter import SemanticChunker
		from langchain_huggingface import HuggingFaceEmbeddings

		embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
		return SemanticChunker(embeddings)

	raise ValueError(
		f"Unknown chunking strategy '{strategy}'. "
		f"Choose one of: {', '.join(CHUNKING_STRATEGIES.values())}"
	)


def _split_document(
	text: str,
	strategy: str,
	chunk_size: int,
	chunk_overlap: int,
) -> list[str]:
	splitter = _create_splitter(strategy, chunk_size, chunk_overlap)

	if strategy == "markdown_header":
		documents = splitter.split_text(text)
		return [document.page_content for document in documents]

	return splitter.split_text(text)


def load_documents(
	file_paths: list[str | Path],
	chunk_size: int = 1000,
	chunk_overlap: int = 150,
	chunking_strategy: str = "recursive_character",
) -> tuple[list[str], list[dict[str, Any]], list[str]]:
	"""Read files and return chunks, metadata, and stable Chroma IDs."""
	chunks: list[str] = []
	metadatas: list[dict[str, Any]] = []
	ids: list[str] = []

	for file_path in file_paths:
		path = Path(file_path)
		text = extract_text(path)
		document_chunks = _split_document(
			text,
			chunking_strategy,
			chunk_size,
			chunk_overlap,
		)
		document_hash = hashlib.sha256(path.read_bytes()).hexdigest()
		source_id = document_hash[:12]

		for chunk_index, chunk in enumerate(document_chunks):
			chunks.append(chunk)
			metadatas.append(
				{
					"source": path.name,
					"file_type": path.suffix.lower(),
					"document_hash": document_hash,
					"chunk_index": chunk_index,
					"chunking_strategy": chunking_strategy,
				}
			)
			ids.append(f"{source_id}-{chunk_index}")

	return chunks, metadatas, ids
